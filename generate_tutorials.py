import sys
from openai import OpenAI
from docx import Document

class MarkdownFile:
    def __init__(self, filename):
        """
        Initializes the MarkdownFile object.
        :param filename: The name of the markdown file.
        """
        if not filename.endswith(".md"):
            filename += ".md"
        self.filename = filename
        self.content = []

    def add_content(self, markdown_content):
        """
        Appends content to the markdown file buffer.
        :param markdown_content: Markdown-formatted string to append.
        """
        self.content.append(markdown_content)

    def save(self):
        """
        Saves the accumulated content to the markdown file.
        """
        try:
            with open(self.filename, "w", encoding="utf-8") as md_file:
                md_file.write("\n".join(self.content))
            print(f"Markdown file '{self.filename}' saved successfully!")
        except Exception as e:
            print(f"An error occurred while saving the file: {e}")

client = OpenAI(api_key="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")

if len(sys.argv) != 3:
    print("Usage: python generate_article.py <article_title> <post_urls>")
    sys.exit(1)

article_title = sys.argv[1]
links_string = sys.argv[2]

output_file = MarkdownFile("input")

def generate_response(prompt):
    try:
        response = client.chat.completions.create(model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": prompt},
        ],
        max_tokens=10000,
        temperature=0.7)
        
        content = response.choices[0].message.content
        print("Response generated successfully!")        
        return content
    
    except Exception as e:
        print(f"An error occurred: {e}")

# Generate Headers
generate_headers = """I want to write a technical tutorial on 'article_header'. 
It will be a purely technical article with code examples. It will be a solution article to a problem related to data structures and algorithms.
I need you to suggest H2 headers only.
I want you to include only headers related to all the different approaches and solutions to the problem in 3 programming languages - Java, Python, C++.
Never use comma in any of the headers.
Include atmost 8-10 headers only, include headers that are more practical and most relevant to the topic.
I want you to make the header SEO optimised based on the keyword from 'article_header'.
Make sure to explain any code examples anywhere.
Please be sure that you include only the most relevant and necessary headers for the tutorial.
The first H2 header must be a modified version of 'article_header'.
The last H2 header should be exactly called 'Frequently Asked Questions'.
All the headers must be very technical and to the point and specific to 'article_header'.
Finally, make sure that you format your output as a list of headers in a string separated by commas. 
For example, the output can be '.....,'Frequently Asked Questions'.
"""

aggregated_header = generate_response(generate_headers.replace("article_header", article_title))

article_headers = [item.strip() for item in aggregated_header.split(",")]

print("-"*20)
print(f"The article headers are: ")
for header in article_headers:
    print(header)
print("-"*20)
print()

# Generate Content and Humanize
introduction_prompt = """
Here are the set of instructions that you strictly have to follow.
1. The article topic is: 'article_title' and the H2 headers in this article are 'article_headers'. 
2. Write two short paragraphs of 3-4 sentences each, the first one should start with the direct definition and solution of the topic. It should not contain any unnecessary things.
3. The second one should contain what we will discuss in this article.
4. Don't include any header or subheader tag in this or any kind of titles. I just want paragraph.
5. Also add a bulleted list of all headers that we will discuss.
6. Make it SEO optimised and follow all SEO best practices. Include meaningful variations of the title keyword multiple times to make it SEO optimised. 
7. Here is a comma separated list of links of few related articles: inner_links. Only use these links in the article wherever possible with proper link and SEO optimised descriptive text. 
8. You should not use any other inner link from any other domain apart from the ones that I have provided you. Don't stuff them or overuse them. You should use a particular link only at max 1 time.
9. If there's code, write it in proper format and HTML.
""".replace("article_title", article_title).replace("article_headers", aggregated_header).replace("inner_links", links_string)

writer_prompt = """
Here are the set of instructions that you strictly have to follow.
1. The article topic is: 'article_title' and the H2 headers in this article are 'article_headers'. 
2. For each header, write relevant content with code if required in short and precise manner. Don't add unnecessary explanations.
4. Make it precise, relevant, and technical. Only if it's necessary, use bullet points only whenever required, if not, just write in paragraphs. 
5. Include code, properties, configurations, examples, wherever applicable. 
6. Maintain structure, formatting, indentation, code blocks, etc. 
7. Write only what is needed in the header. Don't include any unnecessary subheaders. 
8. Here is a comma separated list of links of few related articles: inner_links. Only use these links in the article wherever possible with proper link and SEO optimised descriptive text.
9. You should not use any other inner link from any other domain apart from the ones that I have provided you. Don't stuff them or overuse them. You should use a particular link only at max 1 time.
10. Make it SEO optimised. Follow all SEO best practices. Include meaningful variations of the title keyword multiple times to make it SEO optimised.
11. Strictly follow this: The content that you generate finally should start with the same H2 header that I give you.
12. If there's code, write it in proper format and HTML.
13. Also, the content should not have any conclusion, summary paragraphs or subheaders in it.
14. Heres the H2 header that you have to write for: 
""".replace("article_title", article_title).replace("article_headers", aggregated_header).replace("inner_links", links_string)

faq_prompt = """
1. The article topic is: 'article_title' and the H2 headers are 'article_headers'. 
2. Write a H2 'Frequently Asked Questions' for this article and write 5 technical FAQs that are closely related to this article and are most frequently searched on google. 
3. Each FAQ should be 1 paragraph only atmost 70 words. Make it SEO optimised. 
4. Here is a comma separated list of links of few related articles: inner_links. Only use these links in the article wherever possible with proper link and SEO optimised descriptive text.
5. You should not use any other inner link from any other domain apart from the ones that I have provided you. Don't stuff them or overuse them. You should use a particular link only at max 1 time.
6. Include meaningful variations of the title keyword multiple times to make it SEO optimised. 
7. If there's code, write it in proper format and HTML.
8. Follow all SEO best practices.""".replace("article_title", article_title).replace("article_headers", aggregated_header).replace("inner_links", links_string)

humanizer_prompt = """Rewrite the piece of content in such a way that: 
1. Only use the "we" tone whenever it makes sense. 
2. Replace complex words with easier ones. 
3. Break longer sentences and paragraphs into shorter ones. 
4. Use active voice instead of passive voice. 
5. Use lesser punctuations. 
6. Knowingly, make some small grammatical or punctuational mistakes, not major ones that impacts the SEO.
7. Please write it in such a way that it has been written by a non-native English speaker, who is a beginner technical content writer. 
8. Avoid changing following things - code snippets, technical and Computer science related words like, DevOps, etc., titles and headers. Heres your content:"""

print(f"Generating Response for Introduction Header: {article_headers[0]}")
generated_text = generate_response(introduction_prompt)
print(f"Humanizing Response for Introduction Header: {article_headers[0]}")
humanized_text = generate_response(humanizer_prompt + generated_text)
# doc.add_paragraph(humanized_text)
output_file.add_content(humanized_text + "\n")

for header in article_headers[1:-1]:
    print(f"Generating Response for Header: {header}")
    generated_text = generate_response(writer_prompt + header)
    print(f"Humanizing Response for Header: {header}")
    humanized_text = generate_response(humanizer_prompt + generated_text)
    # doc.add_paragraph(humanized_text)
    output_file.add_content(humanized_text + "\n")

print(f"Generating Response for FAQ Header: {article_headers[-1]}")
generated_text = generate_response(faq_prompt)
print(f"Humanizing Response for FAQ Header: {article_headers[-1]}")
humanized_text = generate_response(humanizer_prompt + generated_text)
# doc.add_paragraph(humanized_text)
output_file.add_content(humanized_text + "\n")

# doc.save(output_file)
output_file.save()